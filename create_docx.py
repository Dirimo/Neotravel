import os
import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('DOCUMENTATION DE PASSATION - PROJET NEOTRAVEL', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle = doc.add_paragraph('Groupe 42')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('1. Présentation du projet', level=1)

doc.add_heading('1.1. Contexte et objectifs', level=2)
doc.add_paragraph(
    "Le projet NeoTravel a pour but de fluidifier la génération de devis pour les transferts en autocar. "
    "L'objectif était de remplacer les formulaires lents par un agent conversationnel (chatbot) intelligent capable de :"
)
p = doc.add_paragraph(style='List Bullet')
p.add_run("Comprendre la demande du client (villes, dates, passagers, type de trajet).")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Interroger notre propre algorithme de tarification.")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Envoyer le devis finalisé dans le chat.")
doc.add_paragraph(
    "En plus de ça, on a ajouté un système de relance par email automatisé via n8n. "
    "L'idée c'est que si un client reçoit son devis mais ne confirme pas la réservation, n8n s'occupe de lui envoyer un mail au bout de quelques jours pour le relancer, ce qui nous aide à améliorer le taux de transformation."
)

doc.add_heading('1.2. Architecture globale', level=2)
doc.add_paragraph("On a mis en place une architecture avec plusieurs briques qui communiquent entre elles :")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Le Frontend (Vue.js / Nuxt 3) : ").bold = True
p.add_run("C'est l'interface client (le composant chat et l'affichage des jolies cartes devis).")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Le Backend intermédiaire (API Nuxt) : ").bold = True
p.add_run("Dans server/api/chat.post.ts, c'est un proxy qui sécurise nos appels vers n8n. Il s'occupe aussi de parser la réponse du bot (avec des Regex) pour générer la petite \"carte devis\" côté front si l'IA donne un prix TTC.")
p = doc.add_paragraph(style='List Bullet')
p.add_run("L'API Métier (Node.js/Express) : ").bold = True
p.add_run("Située dans le dossier src/, elle tourne sur le port 3001. C'est elle qui contient la vraie logique de calcul (le fichier calculer_devis.ts).")
p = doc.add_paragraph(style='List Bullet')
p.add_run("L'Orchestrateur (n8n) : ").bold = True
p.add_run("Il fait tourner notre agent LangChain. Il fait le lien entre le prompt système (system-prompt.txt), la mémoire de la conversation, et les appels HTTP vers notre API métier. Il gère aussi le workflow secondaire pour la relance email.")
p = doc.add_paragraph(style='List Bullet')
p.add_run("La Base de données (Airtable) : ").bold = True
p.add_run("Connectée à n8n pour stocker les informations des devis et gérer le statut de la relance email.")

doc.add_heading('1.3. Stack technique', level=2)
doc.add_paragraph("Frontend : Nuxt 3, Vue 3, TypeScript. Design fait en CSS/Tailwind (dans assets/css/main.css).", style='List Bullet')
doc.add_paragraph("Backend / Microservices : Node.js, Express (pour l'API de devis), TypeScript.", style='List Bullet')
doc.add_paragraph("IA & Workflows : n8n (avec l'intégration LangChain), l'API d'OpenAI (passant par la Vercel AI Gateway pour la sécurité/monitoring).", style='List Bullet')
doc.add_paragraph("Base de données : Airtable.", style='List Bullet')


doc.add_heading('2. Procédure Repreneur', level=1)

doc.add_heading('2.1. Installation et configuration de l’environnement', level=2)
doc.add_paragraph("Pour lancer le projet sur votre machine, voici les étapes qu'on a suivies :")
doc.add_paragraph("1. Cloner le repo et installer les paquets avec `npm install`.")
doc.add_paragraph("2. Dans le dossier, vérifiez les fichiers `.env` si besoin (bien vérifier les ports).")
doc.add_paragraph("3. Lancer le microservice de calcul : Il faut lancer le script src/server.ts (qui tourne sur le port 3001). C'est hyper important sinon le bot ne pourra pas calculer les prix (Timeout).")
doc.add_paragraph("4. Lancer le frontend Nuxt : avec la commande `npm run dev`.")
doc.add_paragraph("5. Configuration n8n :")
doc.add_paragraph("   - Importez le workflow.json (situé dans le dossier n8n) dans votre instance n8n.")
doc.add_paragraph("   - Pensez bien à configurer l'URL de votre n8n dans chat.post.ts (actuellement on a mis http://localhost:5678/webhook-test/neotravel-chat-webhook-001/chat pour que le bouton \"Listen for test event\" de n8n fonctionne).")
doc.add_paragraph("   - Importez également le workflow pour la relance email. Configurez vos credentials (OpenAI, compte Gmail/SMTP, clé API Airtable).")

doc.add_heading('2.2. Interface', level=2)
doc.add_paragraph("Le cœur de l'UI se trouve dans pages/chat.vue. On a utilisé des `ref` pour la réactivité, on y gère le `sessionId` (qui est envoyé à n8n pour garder le contexte du chat) et un affichage conditionnel pour montrer une \"carte\" esthétique si l'API détecte un devis avec prix TTC. Il y a aussi des quick replies (boutons rapides) pour faciliter l'interaction client.")

doc.add_heading('2.3. Évolution du backend', level=2)
doc.add_paragraph("L'API de devis (src/calculer_devis.ts) a été codée avec toutes les règles de NeoTravel : le forfait de base jusqu'à 180km, la majoration au km au-delà, les multiplicateurs selon la saison (basse, moyenne, haute) et le nombre de passagers (tranches de 20-53, 54-63, etc.), plus la marge de 15%.")
doc.add_paragraph("Si vous voulez changer les prix, c'est ce fichier qu'il faut modifier, pas le prompt de l'IA. L'IA ne fait que transmettre les paramètres à cette API, ce qui évite qu'elle \"hallucine\" des prix.")
doc.add_paragraph("Pour la partie relance email, comme tout est géré visuellement dans n8n, c'est super flexible. Le workflow récupère les devis d'Airtable qui ont le statut \"En attente\" depuis plus de 2 jours, génère un petit texte sympa et envoie l'email. Si on veut modifier le délai ou le texte, il suffit de changer les paramètres du nœud dans l'interface n8n sans toucher au code Nuxt.")

doc.add_heading('3. Procédure Équipes Neotravel', level=1)

doc.add_heading('3.1. Accès de la plateforme', level=2)
doc.add_paragraph("Pour les équipes de NeoTravel, l'utilisation est très simple. Ils n'ont pas accès au code, ils se connectent directement sur leur base Airtable. On a mis en place une vue Dashboard qui liste en temps réel tous les devis générés via le chat.")

doc.add_heading('3.2. Traiter et valider les devis', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run("Par défaut : ").bold = True
p.add_run("Lorsqu'un devis est calculé par l'IA, il apparaît dans Airtable avec le statut \"Nouveau\". L'équipe n'a rien à faire. Le script de relance n8n passera automatiquement le statut à \"Relancé\" après avoir envoyé le mail de suivi au client.")
p = doc.add_paragraph(style='List Bullet')
p.add_run("Validation manuelle : ").bold = True
p.add_run("Si un client appelle l'agence ou répond positivement à l'email, un membre de l'équipe modifie le statut dans Airtable pour le passer à \"Validé\". Ça empêche le workflow n8n d'envoyer d'autres relances inutiles.")

doc.add_heading('3.3. Cas complexe et dashboard', level=2)
doc.add_paragraph("Dans notre code calculer_devis.ts, on a défini que s'il y a plus de 85 passagers, c'est un \"cas complexe\". L'IA informe le client qu'un conseiller va le rappeler au lieu de donner un prix erroné. Dans Airtable, ces dossiers remontent en rouge et ne déclenchent aucune relance automatique. C'est le rôle de l'équipe commerciale de prendre le relais au téléphone pour faire une étude sur mesure.")


doc.add_heading('4. Backlog Priorisé', level=1)

doc.add_heading('4.1. Fonctionnalités critiques', level=2)
doc.add_paragraph("Ce qu'on aimerait faire pour la prochaine version (en priorité) :")
doc.add_paragraph("Intégration d'un paiement en ligne (Stripe) : Le but serait de mettre un lien directement dans la carte de devis du chat (et dans l'email de relance) pour que le client puisse payer un acompte direct, ça fluidifierait énormément les ventes.", style='List Bullet')
doc.add_paragraph("Vérification d'email : Gérer les fausses adresses mail. Si la relance échoue (Hard Bounce), il faut que n8n mette à jour Airtable avec le statut \"Erreur Email\" pour que l'équipe commerciale le voie.", style='List Bullet')

doc.add_heading('4.2. Améliorations importantes', level=2)
doc.add_paragraph("Mise en page PDF : L'email de relance contient le récapitulatif du trajet, mais ça ferait beaucoup plus pro si n8n générait un vrai devis en PDF et le mettait en pièce jointe.", style='List Bullet')
doc.add_paragraph("Sauvegarde de l'historique : Actuellement le sessionId est regénéré quand on rafraîchit la page (dans chat.vue). Il faudrait le stocker dans le localStorage pour que l'utilisateur retrouve son chat s'il revient plus tard.", style='List Bullet')

doc.add_heading('4.3. Évolutions futures', level=2)
doc.add_paragraph("Connecter à l'outil de dispatching : Une fois le devis validé, créer automatiquement la course dans le logiciel de gestion de flotte des autocars.", style='List Bullet')
doc.add_paragraph("Relance multi-canal : Si l'email n'est pas ouvert, envoyer un SMS de relance automatique (en utilisant le nœud Twilio dans n8n).", style='List Bullet')

doc.save(r'c:\Users\yanis\Desktop\neoTravel\NEOTRAVEL - LIVRABLE 3_FINAL.docx')
print("DOCX created.")
